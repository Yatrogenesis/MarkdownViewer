#!/usr/bin/env python3
"""
MarkdownViewer - Visor y editor Markdown ligero con exportación PDF/DOC/DOCX
Aplicación de escritorio persistente con vista dual (editor/preview)
"""

import sys
import os
from pathlib import Path
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QTextEdit, QSplitter, QPushButton, QFileDialog, QMenuBar,
    QMenu, QToolBar, QMessageBox, QFontDialog
)
from PyQt6.QtCore import Qt, QTimer
from PyQt6.QtGui import QAction, QIcon, QFont, QTextDocument
from PyQt6.QtWebEngineWidgets import QWebEngineView
import markdown
from markdown.extensions import fenced_code, tables, toc, nl2br, sane_lists
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

# ReportLab para exportación PDF nativa
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, Preformatted
from reportlab.lib import colors
from reportlab.pdfgen import canvas


class MarkdownViewer(QMainWindow):
    def __init__(self):
        super().__init__()
        self.current_file = None
        self.is_modified = False
        self.edit_mode = True
        self.init_ui()
        self.setup_auto_save()

    def init_ui(self):
        """Inicializar interfaz de usuario"""
        self.setWindowTitle("Markdown Viewer & Editor")
        self.setGeometry(100, 100, 1400, 800)

        # Widget central
        central_widget = QWidget()
        self.setCentralWidget(central_widget)

        # Layout principal
        main_layout = QVBoxLayout(central_widget)

        # Crear menú
        self.create_menu()

        # Crear toolbar
        self.create_toolbar()

        # Splitter para editor y preview
        self.splitter = QSplitter(Qt.Orientation.Horizontal)

        # Editor de texto
        self.editor = QTextEdit()
        self.editor.setFont(QFont("Consolas", 11))
        self.editor.textChanged.connect(self.on_text_changed)
        self.editor.setPlaceholderText("Escribe tu Markdown aquí...")

        # Preview web
        self.preview = QWebEngineView()

        # Agregar al splitter
        self.splitter.addWidget(self.editor)
        self.splitter.addWidget(self.preview)
        self.splitter.setSizes([700, 700])

        main_layout.addWidget(self.splitter)

        # Timer para actualización automática del preview
        self.update_timer = QTimer()
        self.update_timer.timeout.connect(self.update_preview)
        self.update_timer.setSingleShot(True)

        # Actualizar preview inicial
        self.update_preview()

    def create_menu(self):
        """Crear barra de menú"""
        menubar = self.menuBar()

        # Menú Archivo
        file_menu = menubar.addMenu("&Archivo")

        new_action = QAction("&Nuevo", self)
        new_action.setShortcut("Ctrl+N")
        new_action.triggered.connect(self.new_file)
        file_menu.addAction(new_action)

        open_action = QAction("&Abrir...", self)
        open_action.setShortcut("Ctrl+O")
        open_action.triggered.connect(self.open_file)
        file_menu.addAction(open_action)

        save_action = QAction("&Guardar", self)
        save_action.setShortcut("Ctrl+S")
        save_action.triggered.connect(self.save_file)
        file_menu.addAction(save_action)

        save_as_action = QAction("Guardar &como...", self)
        save_as_action.setShortcut("Ctrl+Shift+S")
        save_as_action.triggered.connect(self.save_file_as)
        file_menu.addAction(save_as_action)

        file_menu.addSeparator()

        # Submenú Exportar
        export_menu = file_menu.addMenu("&Exportar")

        export_pdf_action = QAction("Exportar a &PDF", self)
        export_pdf_action.triggered.connect(self.export_to_pdf)
        export_menu.addAction(export_pdf_action)

        export_docx_action = QAction("Exportar a &DOCX", self)
        export_docx_action.triggered.connect(self.export_to_docx)
        export_menu.addAction(export_docx_action)

        export_html_action = QAction("Exportar a &HTML", self)
        export_html_action.triggered.connect(self.export_to_html)
        export_menu.addAction(export_html_action)

        file_menu.addSeparator()

        exit_action = QAction("&Salir", self)
        exit_action.setShortcut("Ctrl+Q")
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)

        # Menú Edición
        edit_menu = menubar.addMenu("&Edición")

        undo_action = QAction("&Deshacer", self)
        undo_action.setShortcut("Ctrl+Z")
        undo_action.triggered.connect(self.editor.undo)
        edit_menu.addAction(undo_action)

        redo_action = QAction("&Rehacer", self)
        redo_action.setShortcut("Ctrl+Y")
        redo_action.triggered.connect(self.editor.redo)
        edit_menu.addAction(redo_action)

        edit_menu.addSeparator()

        font_action = QAction("&Fuente...", self)
        font_action.triggered.connect(self.change_font)
        edit_menu.addAction(font_action)

        # Menú Formato
        format_menu = menubar.addMenu("F&ormato")

        # Encabezados
        header_menu = format_menu.addMenu("&Encabezados")
        for i in range(1, 7):
            action = QAction(f"Encabezado {i} (H{i})", self)
            action.triggered.connect(lambda checked, level=i: self.insert_header(level))
            header_menu.addAction(action)

        format_menu.addSeparator()

        bold_action = QAction("&Negrita", self)
        bold_action.setShortcut("Ctrl+B")
        bold_action.triggered.connect(lambda: self.insert_format("**", "**"))
        format_menu.addAction(bold_action)

        italic_action = QAction("&Cursiva", self)
        italic_action.setShortcut("Ctrl+I")
        italic_action.triggered.connect(lambda: self.insert_format("*", "*"))
        format_menu.addAction(italic_action)

        code_action = QAction("&Código inline", self)
        code_action.setShortcut("Ctrl+K")
        code_action.triggered.connect(lambda: self.insert_format("`", "`"))
        format_menu.addAction(code_action)

        format_menu.addSeparator()

        list_action = QAction("Lista &desordenada", self)
        list_action.triggered.connect(self.insert_unordered_list)
        format_menu.addAction(list_action)

        ordered_list_action = QAction("Lista &ordenada", self)
        ordered_list_action.triggered.connect(self.insert_ordered_list)
        format_menu.addAction(ordered_list_action)

        format_menu.addSeparator()

        link_action = QAction("Insertar &enlace", self)
        link_action.setShortcut("Ctrl+L")
        link_action.triggered.connect(self.insert_link)
        format_menu.addAction(link_action)

        image_action = QAction("Insertar i&magen", self)
        image_action.triggered.connect(self.insert_image)
        format_menu.addAction(image_action)

        table_action = QAction("Insertar &tabla", self)
        table_action.triggered.connect(self.insert_table)
        format_menu.addAction(table_action)

        codeblock_action = QAction("Bloque de código", self)
        codeblock_action.triggered.connect(self.insert_code_block)
        format_menu.addAction(codeblock_action)

        # Menú Vista
        view_menu = menubar.addMenu("&Vista")

        toggle_mode_action = QAction("Alternar &modo edición/vista", self)
        toggle_mode_action.setShortcut("F5")
        toggle_mode_action.triggered.connect(self.toggle_edit_mode)
        view_menu.addAction(toggle_mode_action)

        split_view_action = QAction("Vista &dividida", self)
        split_view_action.setShortcut("F6")
        split_view_action.triggered.connect(self.set_split_view)
        view_menu.addAction(split_view_action)

        editor_only_action = QAction("Solo &editor", self)
        editor_only_action.triggered.connect(self.set_editor_only)
        view_menu.addAction(editor_only_action)

        preview_only_action = QAction("Solo &preview", self)
        preview_only_action.triggered.connect(self.set_preview_only)
        view_menu.addAction(preview_only_action)

    def create_toolbar(self):
        """Crear toolbar con botones rápidos"""
        toolbar = QToolBar()
        self.addToolBar(toolbar)

        # Botones de archivo
        new_btn = QPushButton("Nuevo")
        new_btn.clicked.connect(self.new_file)
        toolbar.addWidget(new_btn)

        open_btn = QPushButton("Abrir")
        open_btn.clicked.connect(self.open_file)
        toolbar.addWidget(open_btn)

        save_btn = QPushButton("Guardar")
        save_btn.clicked.connect(self.save_file)
        toolbar.addWidget(save_btn)

        toolbar.addSeparator()

        # Botones de formato
        h1_btn = QPushButton("H1")
        h1_btn.clicked.connect(lambda: self.insert_header(1))
        toolbar.addWidget(h1_btn)

        h2_btn = QPushButton("H2")
        h2_btn.clicked.connect(lambda: self.insert_header(2))
        toolbar.addWidget(h2_btn)

        bold_btn = QPushButton("B")
        bold_btn.setFont(QFont("Arial", 10, QFont.Weight.Bold))
        bold_btn.clicked.connect(lambda: self.insert_format("**", "**"))
        toolbar.addWidget(bold_btn)

        italic_btn = QPushButton("I")
        italic_btn.setFont(QFont("Arial", 10, QFont.Weight.Normal, True))
        italic_btn.clicked.connect(lambda: self.insert_format("*", "*"))
        toolbar.addWidget(italic_btn)

        code_btn = QPushButton("Code")
        code_btn.clicked.connect(lambda: self.insert_format("`", "`"))
        toolbar.addWidget(code_btn)

        toolbar.addSeparator()

        # Botones de vista
        toggle_btn = QPushButton("Vista/Edición")
        toggle_btn.clicked.connect(self.toggle_edit_mode)
        toolbar.addWidget(toggle_btn)

        toolbar.addSeparator()

        # Botones de exportación
        export_pdf_btn = QPushButton("Exportar PDF")
        export_pdf_btn.clicked.connect(self.export_to_pdf)
        toolbar.addWidget(export_pdf_btn)

        export_docx_btn = QPushButton("Exportar DOCX")
        export_docx_btn.clicked.connect(self.export_to_docx)
        toolbar.addWidget(export_docx_btn)

    def on_text_changed(self):
        """Manejar cambios en el texto"""
        self.is_modified = True
        self.update_title()
        # Reiniciar timer para actualización del preview (delay 500ms)
        self.update_timer.start(500)

    def update_preview(self):
        """Actualizar vista previa del Markdown"""
        markdown_text = self.editor.toPlainText()
        html = self.markdown_to_html(markdown_text)
        self.preview.setHtml(html)

    def markdown_to_html(self, markdown_text):
        """Convertir Markdown a HTML con estilo"""
        # Configurar extensiones de Markdown
        md = markdown.Markdown(extensions=[
            'fenced_code',
            'tables',
            'toc',
            'nl2br',
            'sane_lists',
            'codehilite',
            'extra'
        ])

        html_content = md.convert(markdown_text)

        # Envolver con CSS
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{
                    font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                    line-height: 1.6;
                    color: #333;
                    max-width: 900px;
                    margin: 20px auto;
                    padding: 20px;
                    background-color: #f9f9f9;
                }}
                h1, h2, h3, h4, h5, h6 {{
                    color: #2c3e50;
                    margin-top: 24px;
                    margin-bottom: 16px;
                    font-weight: 600;
                    line-height: 1.25;
                }}
                h1 {{
                    font-size: 2em;
                    border-bottom: 2px solid #eaecef;
                    padding-bottom: 0.3em;
                }}
                h2 {{
                    font-size: 1.5em;
                    border-bottom: 1px solid #eaecef;
                    padding-bottom: 0.3em;
                }}
                h3 {{ font-size: 1.25em; }}
                h4 {{ font-size: 1em; }}
                h5 {{ font-size: 0.875em; }}
                h6 {{ font-size: 0.85em; color: #6a737d; }}

                p {{
                    margin-bottom: 16px;
                }}

                code {{
                    background-color: #f6f8fa;
                    padding: 2px 6px;
                    border-radius: 3px;
                    font-family: 'Consolas', 'Monaco', monospace;
                    font-size: 0.9em;
                }}

                pre {{
                    background-color: #f6f8fa;
                    padding: 16px;
                    border-radius: 6px;
                    overflow-x: auto;
                    margin-bottom: 16px;
                }}

                pre code {{
                    background-color: transparent;
                    padding: 0;
                }}

                blockquote {{
                    border-left: 4px solid #dfe2e5;
                    padding-left: 16px;
                    color: #6a737d;
                    margin-left: 0;
                }}

                table {{
                    border-collapse: collapse;
                    width: 100%;
                    margin-bottom: 16px;
                }}

                table th, table td {{
                    border: 1px solid #dfe2e5;
                    padding: 8px 12px;
                    text-align: left;
                }}

                table th {{
                    background-color: #f6f8fa;
                    font-weight: 600;
                }}

                table tr:nth-child(even) {{
                    background-color: #f9f9f9;
                }}

                ul, ol {{
                    padding-left: 2em;
                    margin-bottom: 16px;
                }}

                li {{
                    margin-bottom: 4px;
                }}

                a {{
                    color: #0366d6;
                    text-decoration: none;
                }}

                a:hover {{
                    text-decoration: underline;
                }}

                img {{
                    max-width: 100%;
                    height: auto;
                }}

                hr {{
                    border: 0;
                    border-top: 2px solid #eaecef;
                    margin: 24px 0;
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        return html

    def update_title(self):
        """Actualizar título de la ventana"""
        title = "Markdown Viewer & Editor"
        if self.current_file:
            title = f"{Path(self.current_file).name} - {title}"
        if self.is_modified:
            title = f"*{title}"
        self.setWindowTitle(title)

    def new_file(self):
        """Crear nuevo archivo"""
        if self.check_save_changes():
            self.editor.clear()
            self.current_file = None
            self.is_modified = False
            self.update_title()

    def open_file(self):
        """Abrir archivo Markdown"""
        if self.check_save_changes():
            file_path, _ = QFileDialog.getOpenFileName(
                self,
                "Abrir archivo Markdown",
                "",
                "Archivos Markdown (*.md *.markdown);;Todos los archivos (*.*)"
            )

            if file_path:
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    self.editor.setPlainText(content)
                    self.current_file = file_path
                    self.is_modified = False
                    self.update_title()
                except Exception as e:
                    QMessageBox.critical(self, "Error", f"Error al abrir archivo:\n{str(e)}")

    def save_file(self):
        """Guardar archivo actual"""
        if self.current_file:
            return self.save_to_file(self.current_file)
        else:
            return self.save_file_as()

    def save_file_as(self):
        """Guardar archivo como..."""
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Guardar archivo",
            "",
            "Archivos Markdown (*.md);;Todos los archivos (*.*)"
        )

        if file_path:
            if not file_path.endswith('.md'):
                file_path += '.md'
            return self.save_to_file(file_path)
        return False

    def save_to_file(self, file_path):
        """Guardar contenido a archivo específico"""
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(self.editor.toPlainText())
            self.current_file = file_path
            self.is_modified = False
            self.update_title()
            return True
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error al guardar archivo:\n{str(e)}")
            return False

    def check_save_changes(self):
        """Verificar si hay cambios sin guardar"""
        if self.is_modified:
            reply = QMessageBox.question(
                self,
                "Cambios sin guardar",
                "¿Desea guardar los cambios antes de continuar?",
                QMessageBox.StandardButton.Save |
                QMessageBox.StandardButton.Discard |
                QMessageBox.StandardButton.Cancel
            )

            if reply == QMessageBox.StandardButton.Save:
                return self.save_file()
            elif reply == QMessageBox.StandardButton.Cancel:
                return False
        return True

    def export_to_pdf(self):
        """Exportar a PDF usando ReportLab (nativo)"""
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Exportar a PDF",
            "",
            "Archivos PDF (*.pdf)"
        )

        if file_path:
            if not file_path.endswith('.pdf'):
                file_path += '.pdf'

            try:
                self.export_to_pdf_native(file_path)
                QMessageBox.information(self, "Éxito", f"PDF exportado a:\n{file_path}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Error al exportar PDF:\n{str(e)}")

    def export_to_pdf_native(self, file_path):
        """Exportar a PDF usando ReportLab (sin dependencias externas)"""
        markdown_text = self.editor.toPlainText()

        # Crear documento PDF
        doc = SimpleDocTemplate(
            file_path,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )

        # Estilos
        styles = getSampleStyleSheet()

        # Estilos personalizados
        styles.add(ParagraphStyle(
            name='CustomHeading1',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=12,
            spaceBefore=12
        ))

        styles.add(ParagraphStyle(
            name='CustomHeading2',
            parent=styles['Heading2'],
            fontSize=18,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=10,
            spaceBefore=10
        ))

        styles.add(ParagraphStyle(
            name='CustomHeading3',
            parent=styles['Heading3'],
            fontSize=14,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=8,
            spaceBefore=8
        ))

        styles.add(ParagraphStyle(
            name='CustomCode',
            parent=styles['Code'],
            fontSize=9,
            fontName='Courier',
            backgroundColor=colors.HexColor('#f6f8fa'),
            borderPadding=10,
            leftIndent=10,
            rightIndent=10
        ))

        # Lista de elementos del documento
        story = []

        # Parsear Markdown manualmente
        lines = markdown_text.split('\n')
        i = 0
        in_code_block = False
        code_block_lines = []

        while i < len(lines):
            line = lines[i]

            # Bloques de código
            if line.strip().startswith('```'):
                if not in_code_block:
                    in_code_block = True
                    code_block_lines = []
                else:
                    # Finalizar bloque de código
                    code_text = '\n'.join(code_block_lines)
                    story.append(Preformatted(code_text, styles['CustomCode']))
                    story.append(Spacer(1, 12))
                    in_code_block = False
                i += 1
                continue

            if in_code_block:
                code_block_lines.append(line)
                i += 1
                continue

            # Encabezados
            if line.startswith('# '):
                text = line[2:].strip()
                story.append(Paragraph(text, styles['CustomHeading1']))
                story.append(Spacer(1, 6))
            elif line.startswith('## '):
                text = line[3:].strip()
                story.append(Paragraph(text, styles['CustomHeading2']))
                story.append(Spacer(1, 6))
            elif line.startswith('### '):
                text = line[4:].strip()
                story.append(Paragraph(text, styles['CustomHeading3']))
                story.append(Spacer(1, 6))
            elif line.startswith('#### '):
                text = line[5:].strip()
                story.append(Paragraph(text, styles['Heading4']))
                story.append(Spacer(1, 6))
            elif line.startswith('##### '):
                text = line[6:].strip()
                story.append(Paragraph(text, styles['Heading5']))
                story.append(Spacer(1, 6))
            elif line.startswith('###### '):
                text = line[7:].strip()
                story.append(Paragraph(text, styles['Heading6']))
                story.append(Spacer(1, 6))

            # Listas
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                text = line.strip()[2:]
                text = f"• {text}"
                story.append(Paragraph(text, styles['BodyText']))
            elif re.match(r'^\d+\.\s', line.strip()):
                text = line.strip()
                story.append(Paragraph(text, styles['BodyText']))

            # Código inline (convertir backticks a formato)
            elif '`' in line:
                # Convertir `código` a formato monoespaciado
                text = re.sub(r'`([^`]+)`', r'<font name="Courier" color="#c7254e" backColor="#f9f2f4">\1</font>', line)
                if text.strip():
                    story.append(Paragraph(text, styles['BodyText']))
                    story.append(Spacer(1, 6))

            # Negrita y cursiva
            elif '**' in line or '*' in line:
                # Convertir **negrita** y *cursiva*
                text = re.sub(r'\*\*([^\*]+)\*\*', r'<b>\1</b>', line)
                text = re.sub(r'\*([^\*]+)\*', r'<i>\1</i>', text)
                if text.strip():
                    story.append(Paragraph(text, styles['BodyText']))
                    story.append(Spacer(1, 6))

            # Líneas horizontales
            elif line.strip() in ['---', '***', '___']:
                story.append(Spacer(1, 12))
                from reportlab.platypus import HRFlowable
                story.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
                story.append(Spacer(1, 12))

            # Texto normal
            elif line.strip():
                story.append(Paragraph(line, styles['BodyText']))
                story.append(Spacer(1, 6))

            # Líneas vacías
            else:
                story.append(Spacer(1, 12))

            i += 1

        # Construir PDF
        doc.build(story)

    def export_to_docx(self):
        """Exportar a DOCX"""
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Exportar a DOCX",
            "",
            "Archivos Word (*.docx)"
        )

        if file_path:
            if not file_path.endswith('.docx'):
                file_path += '.docx'

            try:
                doc = Document()
                markdown_text = self.editor.toPlainText()

                # Parsear el Markdown manualmente para DOCX
                lines = markdown_text.split('\n')

                for line in lines:
                    # Encabezados
                    if line.startswith('# '):
                        p = doc.add_heading(line[2:], level=1)
                    elif line.startswith('## '):
                        p = doc.add_heading(line[3:], level=2)
                    elif line.startswith('### '):
                        p = doc.add_heading(line[4:], level=3)
                    elif line.startswith('#### '):
                        p = doc.add_heading(line[5:], level=4)
                    elif line.startswith('##### '):
                        p = doc.add_heading(line[6:], level=5)
                    elif line.startswith('###### '):
                        p = doc.add_heading(line[7:], level=6)
                    # Listas
                    elif line.strip().startswith('- ') or line.strip().startswith('* '):
                        doc.add_paragraph(line.strip()[2:], style='List Bullet')
                    elif re.match(r'^\d+\.\s', line.strip()):
                        text = re.sub(r'^\d+\.\s', '', line.strip())
                        doc.add_paragraph(text, style='List Number')
                    # Bloques de código
                    elif line.strip().startswith('```'):
                        continue
                    # Texto normal
                    elif line.strip():
                        doc.add_paragraph(line)
                    # Líneas vacías
                    else:
                        doc.add_paragraph()

                doc.save(file_path)
                QMessageBox.information(self, "Éxito", f"DOCX exportado a:\n{file_path}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Error al exportar DOCX:\n{str(e)}")

    def export_to_html(self):
        """Exportar a HTML"""
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Exportar a HTML",
            "",
            "Archivos HTML (*.html)"
        )

        if file_path:
            if not file_path.endswith('.html'):
                file_path += '.html'

            try:
                html = self.markdown_to_html(self.editor.toPlainText())
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(html)
                QMessageBox.information(self, "Éxito", f"HTML exportado a:\n{file_path}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Error al exportar HTML:\n{str(e)}")

    def insert_header(self, level):
        """Insertar encabezado"""
        prefix = '#' * level + ' '
        cursor = self.editor.textCursor()
        cursor.insertText(prefix)
        self.editor.setFocus()

    def insert_format(self, start, end):
        """Insertar formato (negrita, cursiva, código)"""
        cursor = self.editor.textCursor()
        selected_text = cursor.selectedText()
        cursor.insertText(f"{start}{selected_text}{end}")
        self.editor.setFocus()

    def insert_unordered_list(self):
        """Insertar lista desordenada"""
        cursor = self.editor.textCursor()
        cursor.insertText("- ")
        self.editor.setFocus()

    def insert_ordered_list(self):
        """Insertar lista ordenada"""
        cursor = self.editor.textCursor()
        cursor.insertText("1. ")
        self.editor.setFocus()

    def insert_link(self):
        """Insertar enlace"""
        cursor = self.editor.textCursor()
        selected_text = cursor.selectedText()
        if selected_text:
            cursor.insertText(f"[{selected_text}](url)")
        else:
            cursor.insertText("[texto del enlace](url)")
        self.editor.setFocus()

    def insert_image(self):
        """Insertar imagen"""
        cursor = self.editor.textCursor()
        cursor.insertText("![descripción](ruta/a/imagen.png)")
        self.editor.setFocus()

    def insert_table(self):
        """Insertar tabla"""
        table = """
| Columna 1 | Columna 2 | Columna 3 |
|-----------|-----------|-----------|
| Fila 1    | Dato      | Dato      |
| Fila 2    | Dato      | Dato      |
"""
        cursor = self.editor.textCursor()
        cursor.insertText(table)
        self.editor.setFocus()

    def insert_code_block(self):
        """Insertar bloque de código"""
        cursor = self.editor.textCursor()
        cursor.insertText("```python\n# Tu código aquí\n```")
        self.editor.setFocus()

    def toggle_edit_mode(self):
        """Alternar entre modo edición y solo vista"""
        self.edit_mode = not self.edit_mode
        if self.edit_mode:
            self.set_split_view()
        else:
            self.set_preview_only()

    def set_split_view(self):
        """Vista dividida (editor + preview)"""
        self.splitter.setSizes([700, 700])
        self.editor.setVisible(True)
        self.preview.setVisible(True)

    def set_editor_only(self):
        """Solo editor"""
        self.editor.setVisible(True)
        self.preview.setVisible(False)

    def set_preview_only(self):
        """Solo preview"""
        self.editor.setVisible(False)
        self.preview.setVisible(True)

    def change_font(self):
        """Cambiar fuente del editor"""
        font, ok = QFontDialog.getFont(self.editor.font(), self)
        if ok:
            self.editor.setFont(font)

    def setup_auto_save(self):
        """Configurar guardado automático"""
        self.auto_save_timer = QTimer()
        self.auto_save_timer.timeout.connect(self.auto_save)
        self.auto_save_timer.start(60000)  # Auto-guardar cada 60 segundos

    def auto_save(self):
        """Guardar automáticamente si hay archivo y modificaciones"""
        if self.current_file and self.is_modified:
            self.save_to_file(self.current_file)

    def closeEvent(self, event):
        """Manejar cierre de ventana"""
        if self.check_save_changes():
            event.accept()
        else:
            event.ignore()


def main():
    app = QApplication(sys.argv)
    app.setApplicationName("Markdown Viewer & Editor")

    viewer = MarkdownViewer()
    viewer.show()

    sys.exit(app.exec())


if __name__ == '__main__':
    main()
